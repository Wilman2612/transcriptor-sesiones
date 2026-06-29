"""
Use case: genera el acta final (txt y docx) imitando el formato oficial de las
actas de concejo de la Municipalidad Distrital de Subtanjalla:

  SESIÓN DE CONCEJO ORDINARIA
  En el salón de actos de la Municipalidad Distrital de Subtanjalla, siendo …

  Hablante: texto de la intervención…
  Otro hablante: …

Sin marcas de tiempo, agrupado por turno de hablante (frases consecutivas del
mismo hablante en un solo párrafo), con el hablante en negrita.
"""
from pathlib import Path

from sqlalchemy.orm import Session as DbSession

from app.infrastructure.persistence.repositories import (
    SqlCorrectionRepository,
    SqlSegmentRepository,
    SqlSessionRepository,
)

MUNICIPALIDAD = "Municipalidad Distrital de Subtanjalla"
_MESES = [
    "", "enero", "febrero", "marzo", "abril", "mayo", "junio",
    "julio", "agosto", "septiembre", "octubre", "noviembre", "diciembre",
]


def _fecha_es(d) -> str:
    return f"{d.day} de {_MESES[d.month]} de {d.year}"


def _segment_text(seg, correction_repo) -> str:
    """Texto final de una frase: reescritura libre > corrección > original."""
    if seg.override_text:
        return seg.override_text
    correction = correction_repo.get_by_segment(seg.id)
    return correction.final_text if correction else seg.original_text


def _group_turns(segments, correction_repo) -> list[tuple[str, str]]:
    """Fusiona frases consecutivas del mismo hablante -> [(hablante, texto), ...]."""
    turns: list[tuple[str, list[str]]] = []
    for seg in segments:
        text = _segment_text(seg, correction_repo).strip()
        if not text:
            continue
        if turns and turns[-1][0] == seg.speaker:
            turns[-1][1].append(text)
        else:
            turns.append((seg.speaker, [text]))
    return [(speaker, " ".join(parts)) for speaker, parts in turns]


def export_transcript(session_id: int, db: DbSession, output_dir: str) -> dict[str, str]:
    session_repo = SqlSessionRepository(db)
    segment_repo = SqlSegmentRepository(db)
    correction_repo = SqlCorrectionRepository(db)

    session = session_repo.get(session_id)
    segments = segment_repo.list_by_session(session_id)
    turns = _group_turns(segments, correction_repo)

    # Resolver "Hablante N" -> nombre/cargo real, si se asignó
    import json

    from app.infrastructure.persistence.models import SessionModel
    m = db.get(SessionModel, session_id)
    names = json.loads(m.speaker_names_json) if m and m.speaker_names_json else {}
    turns = [(names.get(speaker, speaker), text) for speaker, text in turns]

    fecha = _fecha_es(session.date)
    apertura = (
        f"En el salón de actos de la {MUNICIPALIDAD}, siendo el día {fecha}, "
        f"se reunieron los miembros del concejo distrital para llevar a cabo la "
        f"sesión de concejo, bajo el siguiente desarrollo:"
    )

    output_dir = Path(output_dir)
    output_dir.mkdir(parents=True, exist_ok=True)
    safe_name = "".join(c if c.isalnum() or c in "- _" else "_" for c in session.name)

    # ── TXT ──────────────────────────────────────────────────────────────────
    txt_lines = ["SESIÓN DE CONCEJO ORDINARIA", "", apertura, ""]
    txt_lines += [f"{speaker}: {text}" for speaker, text in turns]
    txt_path = output_dir / f"{safe_name}.txt"
    txt_path.write_text("\n\n".join(txt_lines), encoding="utf-8")

    # ── DOCX (formato oficial) ───────────────────────────────────────────────
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(11)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("SESIÓN DE CONCEJO ORDINARIA")
    run.bold = True
    run.font.size = Pt(13)

    doc.add_paragraph("")
    op = doc.add_paragraph(apertura)
    op.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    doc.add_paragraph("")

    for speaker, text in turns:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.add_run(f"{speaker}: ").bold = True
        p.add_run(text)

    docx_path = output_dir / f"{safe_name}.docx"
    doc.save(str(docx_path))

    return {"txt": str(txt_path), "docx": str(docx_path)}
